# metarag/core/loader.py

import os
from typing import List
from pathlib import Path
from abc import ABC, abstractmethod

from dataclasses import dataclass, field

@dataclass
class ExtensionStats:
    count: int
    files: list[str]
    reason: str | None = None


class ExtensionCollection:

    def __init__(self, data: dict):
        self.by_extension = {
            ext: ExtensionStats(**info)
            for ext, info in data.items()
        }

    @property
    def count(self):
        return sum(v.count for v in self.by_extension.values())

    @property
    def files(self):
        files = []
        for stat in self.by_extension.values():
            files.extend(stat.files)
        return files

    def __getitem__(self, ext):
        return self.by_extension[ext]
    def items(self):
        return self.by_extension.items()

    def keys(self):
        return self.by_extension.keys()

    def values(self):
        return self.by_extension.values()

    def __iter__(self):
        return iter(self.by_extension)

    def __len__(self):
        return len(self.by_extension)

    def __contains__(self, key):
        return key in self.by_extension

    def get(self, key, default=None):
        return self.by_extension.get(key, default)
    

class Document:
    """Simple document representation."""

    def __init__(self, text: str, metadata: dict = None):
        self.text = text
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document({len(self.text)} chars, {self.metadata})"


class LoaderInterface(ABC):
    """Contract for document loaders."""

    @abstractmethod
    def load(self) -> List[Document]:
        """Load documents from source."""
        pass


class _MissingDependency(Exception):
    """
    Raised internally when an optional dependency (pypdf, python-docx,
    beautifulsoup4) isn't installed. Caught by _load_file() and aggregated
    into the end-of-load() summary — never printed per-file, so loading
    100 PDFs without pypdf installed prints ONE line, not 100.
    """
    pass


class DocumentList(list):
    """
    A list of Document objects that ALSO carries load/skip stats.

    Behaves exactly like a normal list everywhere (iteration, indexing,
    len(), passing to Chunker.chunk_documents(), isinstance(x, list) checks
    all work unchanged) — it just additionally exposes .loaded and .skipped
    for inspection after load() returns.

    .loaded / .skipped shape (matched, for consistency):
        {
            "pdf": {"count": 8, "files": ["report1.pdf", "manual.pdf", ...]},
            ...
        }
    .skipped entries also include "reason" (why the file was skipped).
    """

    def __init__(self, docs: List[Document], loaded: dict, skipped: dict):
        super().__init__(docs)
        self.loaded = ExtensionCollection(loaded)
        self.skipped = ExtensionCollection(skipped)


class DocumentLoader(LoaderInterface):
    """
    Load documents from various formats.
    No hard dependencies — tries to import only when needed.
    Guides user if format not supported.
    """

    def __init__(self, path: str):
        self.path = Path(path)
        if not self.path.exists():
            raise FileNotFoundError(f"Path not found: {path}")

        # Tracked as we go; used for the printed summary AND exposed on the
        # DocumentList that load() returns. Shape: {ext: {"count": N, "files": [...]}}
        self.loaded: dict = {}
        self.skipped: dict = {}

    def load(self, verbose: bool = True) -> DocumentList:
        """
        Load all documents from path (directory or file).

        Args:
            verbose: print the load summary. Default True — set to False when
                    called from inside another component (like MetaRAG) that
                    already logs its own progress, to avoid printing twice.
        """
        self.loaded = {}
        self.skipped = {}
        docs: List[Document] = []

        if self.path.is_file():
            docs.extend(self._load_file(self.path))
        elif self.path.is_dir():
            for file_path in self.path.rglob("*"):
                if file_path.is_file():
                    docs.extend(self._load_file(file_path))

        if verbose:
            self._print_summary()

        return DocumentList(docs, self.loaded, self.skipped)
    
    # ─────────────────────────────────────────────────────────
    # Dispatch + tracking
    # ─────────────────────────────────────────────────────────

    def _load_file(self, file_path: Path) -> List[Document]:
        """Load a single file based on extension, recording the outcome
        into self.loaded / self.skipped as it goes."""
        ext = file_path.suffix.lower().lstrip(".")

        loader_map = {
            "pdf": self._load_pdf, "txt": self._load_txt, "docx": self._load_docx,
            "html": self._load_html, "htm": self._load_html, "json": self._load_json,
            "csv": self._load_csv, "md": self._load_markdown,
        }
        loader_fn = loader_map.get(ext)
        if loader_fn is None:
            return []  # unsupported format — silent skip, unchanged from before

        try:
            result = loader_fn(file_path)
            if result:
                entry = self.loaded.setdefault(ext, {"count": 0, "files": []})
                entry["count"] += 1
                entry["files"].append(file_path.name)
            return result
        except _MissingDependency as e:
            entry = self.skipped.setdefault(ext, {"count": 0, "files": [], "reason": str(e)})
            entry["count"] += 1
            entry["files"].append(file_path.name)
            return []
        except Exception as e:
            print(f"[DocumentLoader] Error loading {file_path.name}: {e}")
            return []

    def _print_summary(self):
        """One clean report after load() finishes — never printed per-file."""
        print("\nDocumentLoader Report")
        print("-" * 30)

        if self.loaded:
            print("\nFiles Loaded:")
            for ext, info in sorted(self.loaded.items()):
                print(f"  ✓ {ext:<5}: {info['count']}")

        if self.skipped:
            print("\nFiles Skipped:")
            for ext, info in sorted(self.skipped.items()):
                reason_line = info["reason"].splitlines()[-1].strip()
                print(f"  ✗ {ext:<5}: {info['count']} ({reason_line})")

        total_loaded = sum(v["count"] for v in self.loaded.values())
        total_skipped = sum(v["count"] for v in self.skipped.values())
        suffix = f", skipped {total_skipped}" if total_skipped else ""
        print()
        
        print(f"Files Loaded : {total_loaded}")
        print(f"Files Skipped: {total_skipped}")
        print()

    def names(self, which: str = "all", ext: str = None) -> dict:
        """
        Print + return just the filenames, on request.

        Args:
            which: "loaded", "skipped", or "all" (default)
            ext: optionally filter to one extension, e.g. "pdf"
        """
        source = {"loaded": self.loaded, "skipped": self.skipped}
        to_show = source if which == "all" else {which: source[which]}

        result = {}
        for category, by_ext in to_show.items():
            print(f"\n{category.title()}:")
            result[category] = {}
            for extension, info in by_ext.items():
                if ext and extension != ext:
                    continue
                result[category][extension] = info["files"]
                for f in info["files"]:
                    print(f"  {extension}: {f}")

        return result

    # ─────────────────────────────────────────────────────────
    # PDF
    # ─────────────────────────────────────────────────────────

    def _load_pdf(self, file_path: Path) -> List[Document]:
        try:
            import pypdf
        except ImportError:
            raise _MissingDependency(
                "Missing optional dependency 'pypdf'.\n"
                "Use pip to install pypdf: pip install metarag[pdf]"
            )

        docs = []
        try:
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        text=text,
                        metadata={"source": file_path.name, "page": i, "type": "pdf"}
                    ))
        except Exception as e:
            print(f"[DocumentLoader] Error parsing PDF {file_path.name}: {e}")

        return docs

    # ─────────────────────────────────────────────────────────
    # Text
    # ─────────────────────────────────────────────────────────

    def _load_txt(self, file_path: Path) -> List[Document]:
        try:
            text = file_path.read_text(encoding="utf-8")
            if text.strip():
                return [Document(text=text, metadata={"source": file_path.name, "type": "txt"})]
        except UnicodeDecodeError:
            print(f"[DocumentLoader] Could not decode {file_path.name} (not UTF-8)")
        return []

    # ─────────────────────────────────────────────────────────
    # DOCX
    # ─────────────────────────────────────────────────────────

    def _load_docx(self, file_path: Path) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise _MissingDependency(
                "Missing optional dependency 'python-docx'.\n"
                "Use pip to install python-docx: pip install metarag[docx]"
            )

        try:
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if text.strip():
                return [Document(text=text, metadata={"source": file_path.name, "type": "docx"})]
        except Exception as e:
            print(f"[DocumentLoader] Error parsing DOCX {file_path.name}: {e}")

        return []

    # ─────────────────────────────────────────────────────────
    # HTML
    # ─────────────────────────────────────────────────────────

    def _load_html(self, file_path: Path) -> List[Document]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise _MissingDependency(
                "Missing optional dependency 'beautifulsoup4'.\n"
                "Use pip to install beautifulsoup4: pip install metarag[html]"
            )

        try:
            html = file_path.read_text(encoding="utf-8")
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
            if text.strip():
                return [Document(text=text, metadata={"source": file_path.name, "type": "html"})]
        except Exception as e:
            print(f"[DocumentLoader] Error parsing HTML {file_path.name}: {e}")

        return []

    # ─────────────────────────────────────────────────────────
    # JSON
    # ─────────────────────────────────────────────────────────

    def _load_json(self, file_path: Path) -> List[Document]:
        import json
        try:
            data = json.loads(file_path.read_text())
            docs = []
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if isinstance(item, dict) and "text" in item:
                        metadata = item.get("metadata", {})
                        metadata.update({"source": file_path.name, "type": "json", "index": i})
                        docs.append(Document(text=item["text"], metadata=metadata))
            return docs
        except json.JSONDecodeError as e:
            print(f"[DocumentLoader] Invalid JSON in {file_path.name}: {e}")
        except Exception as e:
            print(f"[DocumentLoader] Error parsing JSON {file_path.name}: {e}")
        return []

    # ─────────────────────────────────────────────────────────
    # CSV (pandas is a core dependency — never optional here)
    # ─────────────────────────────────────────────────────────

    def _load_csv(self, file_path: Path) -> List[Document]:
        import pandas as pd
        try:
            df = pd.read_csv(file_path)
            if "text" not in df.columns:
                print(f"[DocumentLoader] CSV must have 'text' column. Found: {list(df.columns)}")
                return []

            docs = []
            for idx, row in df.iterrows():
                metadata = {col: str(row[col]) for col in df.columns if col != "text"}
                metadata.update({"source": file_path.name, "type": "csv", "row": idx})
                docs.append(Document(text=row["text"], metadata=metadata))
            return docs
        except Exception as e:
            print(f"[DocumentLoader] Error parsing CSV {file_path.name}: {e}")
        return []

    # ─────────────────────────────────────────────────────────
    # Markdown
    # ─────────────────────────────────────────────────────────

    def _load_markdown(self, file_path: Path) -> List[Document]:
        try:
            text = file_path.read_text(encoding="utf-8")
            if text.strip():
                return [Document(text=text, metadata={"source": file_path.name, "type": "markdown"})]
        except Exception as e:
            print(f"[DocumentLoader] Error parsing Markdown {file_path.name}: {e}")
        return []

    # ─────────────────────────────────────────────────────────
    # Convenience filters
    # ─────────────────────────────────────────────────────────

    def load_pdfs(self) -> List[Document]:
        return [doc for doc in self.load() if doc.metadata.get("type") == "pdf"]

    def load_texts(self) -> List[Document]:
        return [doc for doc in self.load() if doc.metadata.get("type") == "txt"]

    def load_jsons(self) -> List[Document]:
        return [doc for doc in self.load() if doc.metadata.get("type") == "json"]

    def load_format(self, fmt: str) -> List[Document]:
        return [doc for doc in self.load() if doc.metadata.get("type") == fmt.lower()]